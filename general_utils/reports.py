from typing import Optional, Union

import docx
import math2docx
from docx.shared import Inches

import matplotlib as mpl

import polars as pl


class BaseBlock:
    """
    Базовый класс для элементов документа. Не используется напрямую.
    Реализует общие методы для работы с docx и выравниванием.
    """
    MAP_ALIGMENT = {
        'center': docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER,
        'left': docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT,
        'right': docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT,
        'bottom': docx.enum.table.WD_ALIGN_VERTICAL.BOTTOM,
        'center': docx.enum.table.WD_ALIGN_VERTICAL.CENTER,
        'top': docx.enum.table.WD_ALIGN_VERTICAL.TOP,
    }
    def add_docx(self, dock: docx.Document):
        assert False, 'Abstract method, use instance of class'

    def delete_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

# Base blocks
class PageBreak(BaseBlock):
    """
    Вставляет разрыв страницы в документ.
    
    Пример использования:
    >>> elements.append(PageBreak())  # Добавляет разрыв страницы
    """
    def __init__(self):
        super().__init__()
    
    def add_docx(self, dock: docx.Document):
        dock.add_page_break()

class Paragraph(BaseBlock):
    """
    Создает текстовый абзац с настройками форматирования.
    
    Параметры:
    - text: Текст абзаца
    - style_name: Название стиля (из шаблона docx)
    - is_bold: Жирное начертание
    - is_cursive: Курсивное начертание
    - x_align: Горизонтальное выравнивание ('left', 'center', 'right')
    - y_align: Вертикальное выравнивание ('top', 'center', 'bottom')
    
    Пример использования:
    >>> p = Paragraph(
    ...     "Анализ данных за 2023 год",
    ...     style_name="Heading1",
    ...     is_bold=True,
    ...     x_align='center'
    ... )
    >>> elements.append(p)
    """
    def __init__(
        self,
        text: str,
        style_name: Optional[str] = None,
        is_bold: Optional[bool] = None,
        is_cursive: Optional[bool] = None,
        x_align: Optional[str] = None,
        y_align: Optional[str] = None
    ):
        super().__init__()
        self.text = text
        self.style_name = style_name
        self.is_bold = is_bold
        self.is_cursive = is_cursive
        self.x_align = x_align
        self.y_align = y_align

    def add_docx(self, el: Union[docx.Document, docx.table._Cell]):
        p = el.add_paragraph(self.text, style = self.style_name)
        if self.is_bold is not None: 
            p.runs[0].bold = self.is_bold
        if self.is_cursive is not None: 
            p.runs[0].italic = self.is_cursive
        if self.x_align is not None:
            p.alignment = self.MAP_ALIGMENT[self.x_align]
        if self.y_align is not None:
            p.vertical_alignment = self.MAP_ALIGMENT[self.y_align]

class Table(BaseBlock):
    """
    Вставляет таблицу из DataFrame Polars с аннотацией и настройками форматирования.
    
    Параметры:
    - df: DataFrame Polars
    - style_name: Стиль таблицы
    - x_align: Горизонтальное выравнивание ячеек
    - y_align: Вертикальное выравнивание ячеек
    - annotation_text: Текст примечания под таблицей
    - bold_first_col: Жирное начертание первого столбца
    - is_show_none: Отображать None значения как пустые ячейки
    
    Пример использования:
    >>> df = pl.DataFrame({"Метрика": ["ROI", "CTR"], "Значение": [15.2, 4.7]})
    >>> tbl = Table(
    ...     df,
    ...     annotation_text="Таблица 1: Ключевые метрики",
    ...     bold_first_col=True
    ... )
    >>> elements.append(tbl)
    """
    def __init__(
        self,
        df: pl.DataFrame,
        style_name: Optional[str] = None,
        x_align: str = 'center',
        y_align: str = 'center',
        annotation_text: Optional[str] = None,
        bold_first_col: Optional[bool] = None,
        is_show_none: bool = False
    ):
        super().__init__()
        self.df = df
        self.style_name = style_name
        self.x_align = x_align
        self.y_align = y_align
        self.annotation_text = annotation_text
        self.bold_first_col = bold_first_col
        self.is_show_none = is_show_none

    def add_docx(self, dock: docx.Document):
        if self.annotation_text is not None:
            Paragraph(
                self.annotation_text,
                style_name = self.style_name,
                is_cursive = True,
                x_align = 'right'
            ).add_docx(
                dock
            )
            
        table = dock.add_table(
            rows = self.df.shape[0] + 1, 
            cols = self.df.shape[1], 
            style = 'Table Grid' # add borders
        )

        for j, colname in enumerate(self.df.columns):
            self.delete_paragraph(table.cell(0, j).paragraphs[0]) #.clear() is not working
            
            Paragraph(
                colname,
                style_name = self.style_name,
                is_bold = True,
                x_align = self.x_align
            ).add_docx(
                table.cell(0, j)
            )
            table.cell(0, j).vertical_alignment = self.MAP_ALIGMENT[self.y_align]
            
        for i, row in enumerate(self.df.iter_rows()):
            for j, item in enumerate(row):
                is_bold = self.bold_first_col if self.bold_first_col is not None and j == 0 else None
                text = str(item) if (item is not None or self.is_show_none) else ' ' # preprocess empty cells in df
                
                self.delete_paragraph(table.cell(i + 1, j).paragraphs[0]) #.clear() is not working
                
                Paragraph(
                    text,
                    style_name = self.style_name,
                    is_bold = is_bold,
                    x_align = self.x_align
                ).add_docx(
                    table.cell(i + 1, j)
                )
                table.cell(i + 1, j).vertical_alignment = self.MAP_ALIGMENT[self.y_align]

class Image(BaseBlock):
    """
    Вставляет изображение из matplotlib Figure с аннотацией.
    
    Параметры:
    - fig: Объект Figure matplotlib
    - filename: Путь для сохранения изображения
    - dpi: Разрешение изображения
    - width_inch: Ширина в дюймах
    - height_inch: Высота в дюймах
    - annotation_text: Текст примечания под изображением
    - annotation_style_name: Стиль аннотации
    - is_tight: Обрезка полей изображения
    
    Пример использования:
    >>> fig, ax = plt.subplots()
    >>> ax.plot([1, 2, 3], [10, 20, 15])
    >>> img = Image(
    ...     fig,
    ...     "plot.png",
    ...     width_inch=6.0,
    ...     annotation_text="Рисунок 1: Динамика показателей"
    ... )
    >>> elements.append(img)
    """
    def __init__(
        self,
        fig: mpl.figure.Figure,
        filename: str,
        dpi: int = 100,
        width_inch: Optional[float] = None,
        height_inch: Optional[float] = None,
        annotation_text: Optional[str] = None,
        annotation_style_name: Optional[str] = None,
        is_tight: bool = True
    ):
        super().__init__()
        self.filename = filename
        self.width_inch = width_inch
        self.height_inch = height_inch
        self.annotation_text = annotation_text
        self.annotation_style_name = annotation_style_name
        
        fig.savefig(filename, dpi = dpi, bbox_inches = "tight" if is_tight else None)

    def add_docx(self, dock: docx.Document):
        if self.annotation_text is not None:
            Paragraph(
                self.annotation_text,
                style_name = self.annotation_style_name,
                is_cursive = True,
                x_align = 'right'
            ).add_docx(
                dock
            )
        
        img = dock.add_picture(
            self.filename, 
            width = Inches(self.width_inch) if self.width_inch else None,
            height = Inches(self.height_inch) if self.height_inch else None
        )

class Latex(BaseBlock):
    def __init__(self, latex_text: str, style_name: str):
        self.latex_text = latex_text
        self.style_name = style_name

    def add_docx(self, dock: docx.Document):
        self.p = dock.add_paragraph(style = self.style_name)
        math2docx.add_math(self.p, self.latex_text)

# flat container of elements
class Elements:
    """
    Контейнер для элементов документа. Генерирует итоговый docx файл.
    
    Параметры:
    - docx_name: Путь к шаблону документа (.docx)
    - elements: Список элементов (Paragraph, Table, Image и др.)
    
    Комплексный пример использования:

    >>> from docx_blocks import Paragraph, Table, Image, PageBreak, Elements
    >>> import polars as pl
    >>> import matplotlib.pyplot as plt
    >>> 
    >>> elements = [
    ...    Paragraph("Отчет по анализу данных", style_name="Title", is_bold=True, x_align='center'),
    ...    
    ...    Table(
    ...        pl.DataFrame({"Категория": ["A", "B"], "Сумма": [450, 780]}),
    ...        annotation_text="Таблица 1: Финансовые показатели",
    ...        bold_first_col=True
    ...    ),
    ...    
    ...    Image(
    ...        plt.figure(figsize=(8, 5)),
    ...        "sales_chart.png",
    ...        width_inch=7.0,
    ...        annotation_text="Рисунок 1: Диаграмма продаж"
    ...    ),
    ...    
    ...    PageBreak()
    ...]
    >>>
    >>> doc = Elements("company_template.docx", elements)
    >>> doc.make_docx("financial_report_Q1.docx")
    """
    def __init__(self, docx_name: str, elements: list):
        self.elements = elements
        self.docx = docx.Document(docx_name)

    def get_list_style_names(self):
        return [s.name for s in self.docx.styles]
    
    def make_docx(self, filename: str):
        print('Start generating:')
        
        for n, el in enumerate(self.elements):
            if isinstance(el, str): 
                print(f'[{n}] skip text block')
                continue

            print(f'[{n}] append block {el.__class__.__name__}')
            el.add_docx(self.docx)

        self.docx.save(filename)
        print('Done!')
